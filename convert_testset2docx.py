# This script is used to build a docx file that supports hyperlink, and 
# avoid issues of Chinese URL breakdown. Also includes some jsonl data
# processing and json data filtering. 

import docx ## pip install python-docx
import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from urllib.parse import quote

def encode_url(original_url):
    return quote(original_url, safe='/:?=&')

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    在段落中添加一个超链接。

    :param paragraph: 一个docx.paragraph对象，超链接将被添加到其中。
    :param url: 要链接到的URL字符串。
    :param text: 链接的显示文本。
    :param color: 链接文本的颜色，默认为 "0000FF" (蓝色)。
    :param underline: 是否给链接文本添加下划线，默认为True。
    """
    # 创建超链接关联
    part = paragraph.part
    r_id = part.relate_to(encode_url(url), RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # 创建hyperlink XML元素
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # 创建包含显示文本的run元素
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # 设置超链接文本样式
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    new_run.append(rPr)
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

fin = open('测试集0606.jsonl', encoding="utf-8").readlines()
data = [json.loads(line) for line in fin]

labels = json.load(open('selections_6_12.json', encoding="utf-8"))
disagree_comments = []
for idx, label in labels.items():
    if label == 'disagree':
        disagree_comments.append(int(idx[8:]))

document = docx.Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
document.styles['Normal'].font.size = docx.shared.Pt(12)


for idx, line in enumerate(data):
    if idx in disagree_comments: continue
    document.add_paragraph("第 {} 条测例".format(idx+1))
    document.add_paragraph("原句： " + line['sentence'])
    document.add_paragraph("原批注： " + line['comment'])
    document.add_paragraph("AI修改后批注： " + line['revised_comment'])
    p = document.add_paragraph("源文档超链接： " )
    line['source_title'] = line['source_title'].replace('\\\\', '/')
    line['source_title'] = line['source_title'].replace('\\', '/')
    line['source_title'] = line['source_title'].replace(' ', '_')
    line['source_title'] = line['source_title'].replace('docx', 'html')
    # print(line['source_title'])
    line['source_title'] = "http://115.182.62.193:11450/html/" + line['source_title']
    # print(line['source_title'])
    add_hyperlink(p, url=line['source_title'], text=line['source_title'])

    document.add_paragraph("\n")
    
document.save('测试集-0606.docx')